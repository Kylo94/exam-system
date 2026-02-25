"""测试Word文档提取是否能保留代码缩进"""

from docx import Document
from docx.shared import Inches, Pt
from pathlib import Path

# 创建一个测试Word文档
doc = Document()

# 添加标题
doc.add_heading('Python代码测试', level=1)

# 添加题目
doc.add_paragraph('1. 编写一个Python函数，实现以下功能：')
doc.add_paragraph('1. 输入一个整数列表')
doc.add_paragraph('2. 返回列表中的最大值')
doc.add_paragraph('3. 处理空列表的情况')

# 添加代码（使用不同方法测试）
doc.add_paragraph('示例代码：')

# 方法1：直接添加（会被去除缩进）
para = doc.add_paragraph("def find_max(numbers):\n    if not numbers:\n        return None\n    max_num = numbers[0]\n    for num in numbers[1:]:\n        if num > max_num:\n            max_num = num\n    return max_num")

# 方法2：使用制表符（也会被去除）
# para2 = doc.add_paragraph("\tdef find_max(numbers):\n\t\tif not numbers:\n\t\t\treturn None")

# 方法3：设置段落缩进
for code_line in [
    "def find_max(numbers):",
    "    if not numbers:",
    "        return None",
    "    max_num = numbers[0]",
    "    for num in numbers[1:]:",
    "        if num > max_num:",
    "            max_num = num",
    "    return max_num"
]:
    para = doc.add_paragraph(code_line)
    # 设置左缩进
    from docx.shared import Pt
    if code_line.startswith("    "):
        para.paragraph_format.left_indent = Pt(28)  # 28pt ≈ 1cm = 4 spaces
    elif code_line.startswith("        "):
        para.paragraph_format.left_indent = Pt(56)  # 56pt ≈ 2cm = 8 spaces

# 保存文档
output_path = "/Users/kylo/Documents/Workspace/答题网站/test_code_word.docx"
doc.save(output_path)
print(f"✓ 测试Word文档已创建: {output_path}")

# 现在测试提取
print("\n" + "=" * 60)
print("测试提取文档")
print("=" * 60)

# 直接用python-docx读取
doc2 = Document(output_path)
print(f"\n原始Word文档段落数: {len(doc2.paragraphs)}")
print(f"前5个段落:")
for i, para in enumerate(doc2.paragraphs[:5]):
    print(f"  段落{i}: {repr(para.text[:50])}")

from app.parsers.ai_document_parser import AIDocumentParser

parser = AIDocumentParser()
text, image_info = parser._extract_text_and_images_from_docx(output_path)

print(f"\n提取的文本长度: {len(text)} 字符")
print(f"\n提取的完整文本:")
print("-" * 60)
print(text)
print("-" * 60)

# 检查是否保留了缩进
lines = text.split('\n')
indented_lines = [line for line in lines if line.startswith('    ') or line.startswith('\t')]

print(f"\n统计:")
print(f"  - 总行数: {len(lines)}")
print(f"  - 缩进行数: {len(indented_lines)}")

if len(indented_lines) > 0:
    print("\n✓ 代码缩进保留成功！")
    print("\n缩进的代码行:")
    for line in indented_lines[:5]:
        print(f"  {line[:60]}")
else:
    print("\n✗ 代码缩进丢失，需要改进提取逻辑")

# 清理测试文件
Path(output_path).unlink()
print(f"\n✓ 测试文档已删除")
