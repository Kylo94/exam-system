"""调试Word文档提取问题"""

from docx import Document
from docx.shared import Pt

# 创建一个简单的测试文档
doc = Document()

# 添加普通段落
doc.add_paragraph('Hello World')

# 添加有缩进的段落
para = doc.add_paragraph('    indented text')
para.paragraph_format.left_indent = Pt(28)

# 保存
output_path = "/Users/kylo/Documents/Workspace/答题网站/test_simple.docx"
doc.save(output_path)

# 读取并检查
doc2 = Document(output_path)
for i, para in enumerate(doc2.paragraphs):
    print(f"段落{i}: text={repr(para.text)}, len={len(para.text)}")

# 检查_element的内容
print("\n检查XML内容:")
for i, para in enumerate(doc2.paragraphs):
    print(f"\n段落{i}:")
    print(f"  text: {repr(para.text)}")
    print(f"  text length: {len(para.text)}")
    print(f"  _element: {para._element.tag}")
    
    # 检查是否有大量空白字符
    if len(para.text) > 100:
        print(f"  WARNING: 文本过长！")
        print(f"  前100字符: {repr(para.text[:100])}")
        print(f"  后100字符: {repr(para.text[-100:])}")

import os
os.unlink(output_path)
