"""文档解析服务"""

from typing import List, Dict, Optional
import os


class DocumentParserService:
    """文档解析服务"""

    def __init__(self, db):
        self.db = db

    def parse_document(self, file_path: str, file_type: str) -> Dict[str, any]:
        """
        解析文档文件

        Args:
            file_path: 文件路径
            file_type: 文件类型 ('docx', 'pdf')

        Returns:
            解析结果，包含：
            - success: 是否成功
            - text: 提取的文本
            - questions: 提取的试题列表
            - summary: 文档摘要
            - error: 错误信息
        """
        result = {
            'success': False,
            'text': '',
            'questions': [],
            'summary': '',
            'error': None
        }

        try:
            # 提取文本
            text = self._extract_text(file_path, file_type)
            result['text'] = text

            # 使用AI提取试题
            from app.services.ai_service import get_ai_service
            ai_service = get_ai_service()

            # 更新最后使用时间
            if hasattr(ai_service, 'provider'):
                from app.models.ai_config import AIConfig
                config = AIConfig.get_active_provider()
                if config:
                    config.update_last_used()

            # 提取试题
            questions = ai_service.extract_questions_from_text(text)
            result['questions'] = questions

            # 生成摘要
            if len(text) > 0:
                summary = ai_service.summarize_document(text)
                result['summary'] = summary

            result['success'] = True

        except Exception as e:
            result['error'] = str(e)

        return result

    def _extract_text(self, file_path: str, file_type: str) -> str:
        """
        从文件中提取文本

        Args:
            file_path: 文件路径
            file_type: 文件类型

        Returns:
            提取的文本
        """
        if file_type == 'docx':
            return self._extract_from_docx(file_path)
        elif file_type == 'pdf':
            return self._extract_from_pdf(file_path)
        else:
            raise Exception(f"不支持的文件类型: {file_type}")

    def _extract_from_docx(self, file_path: str) -> str:
        """
        从Word文档中提取文本

        Args:
            file_path: 文件路径

        Returns:
            提取的文本
        """
        try:
            from docx import Document
            doc = Document(file_path)

            paragraphs = []
            for para in doc.paragraphs:
                paragraphs.append(para.text)

            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.append(cell.text)

            return '\n'.join(paragraphs)

        except ImportError:
            raise Exception("需要安装 python-docx 库: pip install python-docx")
        except Exception as e:
            raise Exception(f"解析Word文档失败: {str(e)}")

    def _extract_from_pdf(self, file_path: str) -> str:
        """
        从PDF文档中提取文本

        Args:
            file_path: 文件路径

        Returns:
            提取的文本
        """
        try:
            import PyPDF2

            text = []
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text.append(page.extract_text())

            return '\n'.join(text)

        except ImportError:
            raise Exception("需要安装 PyPDF2 库: pip install PyPDF2")
        except Exception as e:
            raise Exception(f"解析PDF文档失败: {str(e)}")

    def validate_questions(self, questions: List[Dict]) -> List[Dict]:
        """
        验证提取的试题格式

        Args:
            questions: 试题列表

        Returns:
            验证后的试题列表
        """
        valid_questions = []
        required_fields = ['type', 'content', 'correct_answer']

        valid_types = [
            'single_choice',
            'multiple_choice',
            'judgment',
            'fill_blank',
            'subjective'
        ]

        for idx, question in enumerate(questions):
            # 检查必需字段
            missing_fields = [field for field in required_fields if field not in question]
            if missing_fields:
                print(f"第{idx+1}题缺少字段: {missing_fields}")
                continue

            # 检查题型
            if question['type'] not in valid_types:
                print(f"第{idx+1}题题型无效: {question['type']}")
                continue

            # 检查选项（单选和多选题必须有选项）
            if question['type'] in ['single_choice', 'multiple_choice']:
                if 'options' not in question or not question['options']:
                    print(f"第{idx+1}题缺少选项")
                    continue

            # 设置默认值
            question.setdefault('points', 10)
            question.setdefault('explanation', '')
            question.setdefault('order_index', idx + 1)

            valid_questions.append(question)

        return valid_questions
