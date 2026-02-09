"""Word文档解析器

用于解析包含试题的Word文档。
"""

import re
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path

from .base import BaseParser


class DocxParser(BaseParser):
    """Word文档解析器"""
    
    def __init__(self):
        """初始化Word文档解析器"""
        self._doc = None
        
    def parse(self, file_path: str) -> Dict[str, Any]:
        """解析Word文档
        
        Args:
            file_path: Word文档路径
            
        Returns:
            解析后的文档数据
            
        Raises:
            FileNotFoundError: 文件不存在
            ImportError: python-docx未安装
            ValueError: 解析失败
        """
        try:
            import docx
        except ImportError:
            raise ImportError("请安装python-docx: pip install python-docx")
        
        # 验证文件
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        if path.suffix.lower() not in ['.docx', '.doc']:
            raise ValueError(f"不支持的文件格式: {path.suffix}")
        
        try:
            # 读取文档
            doc = docx.Document(file_path)
            self._doc = doc
            
            # 提取文档信息
            file_info = self.get_file_info(file_path)
            
            # 提取所有段落
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            
            # 提取表格
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            # 提取问题
            questions = self._extract_questions_from_text(paragraphs)
            
            # 构建结果
            result = {
                'file_info': file_info,
                'metadata': {
                    'paragraph_count': len(paragraphs),
                    'table_count': len(tables_data),
                    'question_count': len(questions),
                    'format': 'docx'
                },
                'content': {
                    'paragraphs': paragraphs,
                    'tables': tables_data,
                },
                'questions': questions
            }
            
            return result
            
        except Exception as e:
            raise ValueError(f"文档解析失败: {str(e)}")
    
    def extract_questions(self, content: Dict[str, Any]) -> List[Dict[str, Any]]:
        """从解析内容中提取问题
        
        Args:
            content: 解析后的文档内容
            
        Returns:
            问题字典列表
        """
        return content.get('questions', [])
    
    def validate_question(self, question_data: Dict[str, Any]) -> bool:
        """验证问题数据格式
        
        Args:
            question_data: 问题数据字典
            
        Returns:
            是否有效
        """
        required_fields = ['content', 'type', 'score']
        for field in required_fields:
            if field not in question_data:
                return False
        
        # 验证问题类型
        valid_types = ['single_choice', 'multiple_choice', 'true_false', 'fill_blank', 'short_answer']
        if question_data['type'] not in valid_types:
            return False
        
        # 验证分值
        try:
            score = float(question_data['score'])
            if score <= 0:
                return False
        except (ValueError, TypeError):
            return False
        
        # 根据类型验证额外字段
        q_type = question_data['type']
        
        if q_type in ['single_choice', 'multiple_choice']:
            if 'options' not in question_data:
                return False
            
            options = question_data['options']
            if not isinstance(options, dict) or 'choices' not in options:
                return False
            
            choices = options['choices']
            if not isinstance(choices, list) or len(choices) < 2:
                return False
            
            if 'correct_answer' not in question_data:
                return False
            
            correct_answer = question_data['correct_answer']
            if q_type == 'single_choice' and not isinstance(correct_answer, str):
                return False
            elif q_type == 'multiple_choice' and not isinstance(correct_answer, list):
                return False
        
        elif q_type == 'true_false':
            if 'correct_answer' not in question_data:
                return False
            
            correct_answer = question_data['correct_answer']
            if not isinstance(correct_answer, bool):
                # 也接受字符串形式的布尔值
                if isinstance(correct_answer, str):
                    if correct_answer.lower() not in ['true', 'false']:
                        return False
                else:
                    return False
        
        elif q_type in ['fill_blank', 'short_answer']:
            if 'correct_answer' not in question_data:
                return False
            
            correct_answer = question_data['correct_answer']
            if not isinstance(correct_answer, str):
                return False
        
        return True
    
    def supported_formats(self) -> List[str]:
        """获取支持的文件格式
        
        Returns:
            支持的文件扩展名列表
        """
        return ['.docx', '.doc']
    
    def _extract_questions_from_text(self, paragraphs: List[str]) -> List[Dict[str, Any]]:
        """从段落文本中提取问题
        
        Args:
            paragraphs: 段落文本列表
            
        Returns:
            问题字典列表
        """
        questions = []
        current_question = None
        current_options = []
        in_question = False
        question_number = 1
        
        # 常见问题模式
        question_patterns = [
            r'^\d+[\.、]',  # 1. 或 1、
            r'^\(\d+\)',    # (1)
            r'^第\d+题',    # 第1题
            r'^[一二三四五六七八九十]+[\.、]',  # 一、 二.
        ]
        
        # 选项模式
        option_patterns = [
            r'^[A-D][\.、]',      # A. 或 A、
            r'^[①②③④⑤⑥]',      # ① ②
            r'^\(\w+\)',          # (A) (B)
        ]
        
        for para in paragraphs:
            # 检查是否是问题开始
            is_question = any(re.match(pattern, para) for pattern in question_patterns)
            is_option = any(re.match(pattern, para) for pattern in option_patterns)
            
            if is_question:
                # 保存前一个问题
                if current_question:
                    questions.append(current_question)
                
                # 开始新问题
                question_text = re.sub(r'^\d+[\.、]|^\(\d+\)|^第\d+题|^[一二三四五六七八九十]+[\.、]', '', para).strip()
                
                # 尝试检测问题类型
                q_type = self._detect_question_type(question_text)
                
                current_question = {
                    'number': question_number,
                    'content': question_text,
                    'type': q_type,
                    'score': 1.0,  # 默认分值
                    'options': {'choices': []},
                    'correct_answer': None,
                    'explanation': '',
                    'raw_text': [para]
                }
                
                current_options = []
                in_question = True
                question_number += 1
            
            elif in_question and is_option:
                # 处理选项
                option_text = para.strip()
                current_options.append(option_text)
                current_question['raw_text'].append(para)
                
                # 更新选项列表
                if 'options' in current_question:
                    current_question['options']['choices'] = current_options
            
            elif in_question and para.strip():
                # 可能是问题内容的延续或解析说明
                current_question['raw_text'].append(para)
                
                # 检查是否是答案或解析
                if re.search(r'[答案|正确答案]', para):
                    # 提取答案
                    answer_match = re.search(r'[答案|正确答案][：:]\s*([A-D①②③④]+)', para)
                    if answer_match:
                        answer = answer_match.group(1)
                        # 转换为标准格式
                        if current_question['type'] == 'single_choice':
                            current_question['correct_answer'] = self._normalize_single_answer(answer)
                        elif current_question['type'] == 'multiple_choice':
                            current_question['correct_answer'] = self._normalize_multiple_answer(answer)
                
                elif re.search(r'[解析|说明]', para):
                    current_question['explanation'] = para
        
        # 添加最后一个问题
        if current_question:
            questions.append(current_question)
        
        return questions
    
    def _detect_question_type(self, question_text: str) -> str:
        """检测问题类型
        
        Args:
            question_text: 问题文本
            
        Returns:
            问题类型
        """
        text_lower = question_text.lower()
        
        # 判断题特征
        if re.search(r'(是否正确|对不对|判断.*正误|true.*false|正确.*错误)', text_lower):
            return 'true_false'
        
        # 填空题特征
        if re.search(r'(填空|______|_+|\s+_+\s+)', text_lower):
            return 'fill_blank'
        
        # 简答题特征
        if re.search(r'(简述|论述|分析|说明|解释|为什么|如何|怎样|原因|意义|影响)', text_lower):
            return 'short_answer'
        
        # 多选题特征
        if re.search(r'(多选|哪些|多个|不止一个)', text_lower):
            return 'multiple_choice'
        
        # 默认单选题
        return 'single_choice'
    
    def _normalize_single_answer(self, answer: str) -> str:
        """标准化单选题答案
        
        Args:
            answer: 原始答案
            
        Returns:
            标准化答案（A, B, C, D等）
        """
        if not answer:
            return ''
        
        # 移除空格和标点
        answer = answer.strip().upper()
        
        # 处理中文数字或特殊符号
        mapping = {
            '①': 'A', '②': 'B', '③': 'C', '④': 'D',
            '⑤': 'E', '⑥': 'F', '⑦': 'G', '⑧': 'H',
            '1': 'A', '2': 'B', '3': 'C', '4': 'D',
            '5': 'E', '6': 'F', '7': 'G', '8': 'H',
        }
        
        if answer in mapping:
            return mapping[answer]
        
        # 提取第一个字母
        match = re.search(r'[A-H]', answer)
        if match:
            return match.group(0)
        
        return answer[0] if answer else ''
    
    def _normalize_multiple_answer(self, answer: str) -> List[str]:
        """标准化多选题答案
        
        Args:
            answer: 原始答案
            
        Returns:
            标准化答案列表
        """
        if not answer:
            return []
        
        answer = answer.strip().upper()
        
        # 分割多个答案
        answers = []
        for char in answer:
            if char.isalpha() and 'A' <= char <= 'H':
                answers.append(char)
        
        return list(set(answers))  # 去重
    
    def parse_to_exam(self, file_path: str, exam_title: str, 
                     subject_id: int, level_id: int) -> Dict[str, Any]:
        """解析文档并转换为考试数据
        
        Args:
            file_path: Word文档路径
            exam_title: 考试标题
            subject_id: 科目ID
            level_id: 难度级别ID
            
        Returns:
            考试数据字典
        """
        # 解析文档
        parsed_data = self.parse(file_path)
        
        # 提取问题
        questions = self.extract_questions(parsed_data)
        
        # 验证问题
        valid_questions = []
        for q in questions:
            if self.validate_question(q):
                valid_questions.append(q)
        
        # 构建考试数据
        exam_data = {
            'title': exam_title,
            'subject_id': subject_id,
            'level_id': level_id,
            'description': f"从文档 '{parsed_data['file_info']['name']}' 导入的考试",
            'duration_minutes': len(valid_questions) * 2,  # 预估时间
            'questions': valid_questions,
            'source_file': parsed_data['file_info'],
            'metadata': parsed_data['metadata']
        }
        
        return exam_data