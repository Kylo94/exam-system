"""AI辅助文档解析服务"""

import re
import json
import ast
from typing import Dict, List, Any, Optional
from pathlib import Path

from app.services.ai_service import get_ai_service

# 尝试导入json5库，它支持更宽松的JSON格式
try:
    import json5
    HAS_JSON5 = True
except ImportError:
    HAS_JSON5 = False


class AIDocumentParser:
    """AI辅助的文档解析器"""

    def __init__(self, use_ai: bool = True, progress_callback=None, upload_folder: str = None, extract_images: bool = True):
        """
        初始化解析器

        Args:
            use_ai: 是否使用AI解析，默认True
            progress_callback: 进度回调函数，接收 (percent, message) 参数
            upload_folder: 图片上传文件夹路径，用于存储提取的图片
            extract_images: 是否提取图片，默认True（设为False可跳过图片提取）
        """
        self.use_ai = use_ai
        self.ai_service = None
        self.log_messages = []  # 解析日志记录
        self.progress_callback = progress_callback  # 进度回调函数
        self.upload_folder = upload_folder or 'uploads/images'  # 图片保存目录
        self.extract_images = extract_images  # 是否提取图片
        # 确保图片目录存在（如果需要提取图片）
        if extract_images:
            Path(self.upload_folder).mkdir(parents=True, exist_ok=True)

    def _add_log(self, message: str, level: str = 'info', progress: int = None):
        """添加解析日志"""
        from datetime import datetime
        from flask import current_app

        timestamp = datetime.now().strftime('%H:%M:%S')
        icon = {'info': 'ℹ️', 'success': '✅', 'warning': '⚠️', 'error': '❌'}.get(level, '•')
        log_entry = f"[{timestamp}] {icon} {message}"
        self.log_messages.append(log_entry)

        # 如果有进度百分比或progress_callback，调用进度回调
        if self.progress_callback:
            if progress is not None:
                # 传递进度、消息和级别
                self.progress_callback(progress, message, level)
            else:
                # 只传递消息和级别（不更新进度）
                self.progress_callback(message, level)

        # 同时记录到 Flask 日志
        try:
            if level == 'success':
                current_app.logger.info(message)
            elif level == 'error':
                current_app.logger.error(message)
            elif level == 'warning':
                current_app.logger.warning(message)
            else:
                current_app.logger.debug(message)
        except RuntimeError:
            # 在应用上下文外时只打印
            pass

    def get_logs(self) -> str:
        """获取解析日志的HTML格式"""
        # 过滤掉包含堆栈信息的错误日志（包含"Traceback"或"File \"/"的内容）
        filtered_logs = []
        for log in self.log_messages:
            # 跳过包含详细错误信息的日志
            if 'Traceback' in log or 'File "/opt/' in log or 'File "/Users/' in log:
                continue
            # 跳过太长的错误详情行
            if '  File ' in log and 'line ' in log:
                continue
            filtered_logs.append(log)
        return '<br>'.join(filtered_logs)

    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """
        解析文档，优先使用AI

        Args:
            file_path: 文档路径

        Returns:
            解析结果字典
        """
        self.log_messages = []  # 清空日志
        self._add_log(f"开始解析文档: {Path(file_path).name}", 'info', progress=5)

        # 1. 首先使用规则解析提取文本和图片信息
        self._add_log("步骤1: 从Word文档提取文本内容...", 'info', progress=10)
        try:
            text, image_info = self._extract_text_and_images_from_docx(file_path)
            self._add_log(f"✓ 文本提取完成，共 {len(text)} 个字符", 'success', progress=20)
            self._add_log(f"✓ 检测到 {len(image_info)} 张图片", 'success', progress=25)
        except Exception as e:
            # 图片提取失败，尝试只提取文本
            self._add_log(f"⚠️ 图片提取失败，尝试只提取文本: {str(e)}", 'warning', progress=15)
            try:
                text = self._extract_text_only(file_path)
                image_info = []
                self._add_log(f"✓ 文本提取完成（跳过图片），共 {len(text)} 个字符", 'success', progress=20)
            except Exception as e2:
                # 完全失败
                raise Exception(f"文档提取失败: {str(e2)}")

        # 显示文本内容的前1000字符用于调试
        self._add_log(f"文本预览（前1000字符）:\n{text[:1000]}", 'info')

        for idx, img in enumerate(image_info):
            self._add_log(f"  图片{idx+1}: 位置{img['location']}, 对应题号{img['question_number']}, 类型:{img['image_type']}, 上下文: {img['text_context']}", 'info')

        # 2. 使用AI解析试题
        if self.use_ai:
            self._add_log("步骤2: 使用AI智能解析试题...", 'info', progress=30)
            self._add_log("  正在构建AI提示词...", 'info', progress=35)
            questions = self._parse_with_ai(text, image_info)
        else:
            self._add_log("步骤2: 使用规则解析试题...", 'info', progress=30)
            questions = self._parse_with_rules(text, image_info)

        # 3. 返回结果
        self._add_log(f"步骤3: 解析完成，共识别 {len(questions)} 道试题", 'success', progress=95)

        # 记录每道题的详细信息用于调试
        self._add_log("试题详细信息:", 'info')
        for idx, q in enumerate(questions, 1):
            q_type = q.get('type', 'unknown')
            q_content = q.get('content', '')[:100]  # 只显示前100字符
            q_has_image = q.get('content_has_image', False)
            self._add_log(f"  第{idx}题: 类型={q_type}, 包含图片={q_has_image}, 内容预览: {q_content}", 'info')

        self._add_log("解析成功完成！", 'success', progress=100)

        return {
            'success': True,
            'questions': questions,
            'question_count': len(questions),
            'raw_text': text,
            'parse_method': 'ai' if self.use_ai else 'rules',
            'parse_log': self.get_logs()
        }

    def _extract_text_only(self, file_path: str) -> str:
        """只提取文本，不提取图片（备用方法）"""
        try:
            from docx import Document
            doc = Document(file_path)

            text_parts = []
            for para in doc.paragraphs:
                text_parts.append(para.text)

            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())

            return '\n'.join(text_parts)

        except ImportError:
            raise Exception("需要安装 python-docx: pip install python-docx")
        except Exception as e:
            raise Exception(f"提取文本失败: {str(e)}")

    def _extract_text_and_images_from_docx(self, file_path: str) -> tuple:
        """从Word文档提取文本和图片（保留原始格式）"""
        try:
            from docx import Document
            doc = Document(file_path)

            text_parts = []
            image_info = []
            current_question_number = None
            line_number = 0
            global_image_index = 0  # 全局图片索引，用于生成唯一文件名

            # 首先收集所有段落文本，用于查找上下文
            all_paragraphs_text = []
            for para in doc.paragraphs:
                all_paragraphs_text.append(para.text.strip())

            # 提取段落内容和图片
            for para_idx, para in enumerate(doc.paragraphs):
                line_number += 1

                # 获取段落文本
                para_text = para.text
                stripped_text = para_text.strip()

                # 只添加非空段落或空行
                if para_text or (para_idx > 0 and para.text == '' and len(all_paragraphs_text) > 0):
                    if stripped_text:
                        # 检测题号
                        match = re.match(r'^(\d+)[\.、]\s*', stripped_text)
                        if match:
                            current_question_number = int(match.group(1))

                        text_parts.append(para_text)
                    else:
                        # 保留空行（用于代码块的分隔）
                        text_parts.append('')

                # 提取段落中的图片（如果启用图片提取）
                if self.extract_images:
                    try:
                        for run in para.runs:
                            # 检查run中的drawing元素（嵌入式图片）
                            for drawing in run._element.xpath('.//w:drawing'):
                                # 找到图片
                                pics = drawing.xpath('.//pic:pic')
                                if pics:
                                    # 获取图片的上下文文本（前后段落）
                                    context_text = self._get_image_context(all_paragraphs_text, para_idx, current_question_number)

                                    # 判断图片是在题目中还是选项中
                                    image_type = self._determine_image_type(context_text, current_question_number)

                                    # 提取并保存图片
                                    image_path = self._extract_and_save_image(drawing, doc, global_image_index)
                                    global_image_index += 1

                                    image_info.append({
                                        'location': f'paragraph_{line_number}',
                                        'question_number': current_question_number,
                                        'image_type': image_type,
                                        'text_context': context_text,
                                        'image_path': image_path
                                    })
                    except Exception as e:
                        # 单个图片提取失败不影响整体
                        self._add_log(f"  ⚠️ 段落{para_idx}图片提取失败: {str(e)}", 'warning')

            # 也提取表格内容
            table_line_number = line_number
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            text_parts.append(cell_text)
                            table_line_number += 1

                        # 检查单元格中的图片（如果启用图片提取）
                        if self.extract_images:
                            try:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        for drawing in run._element.xpath('.//w:drawing'):
                                            pics = drawing.xpath('.//pic:pic')
                                            if pics:
                                                context_text = cell.text.strip()
                                                image_type = self._determine_image_type(context_text, current_question_number)

                                                # 提取并保存图片
                                                image_path = self._extract_and_save_image(drawing, doc, global_image_index)
                                                global_image_index += 1

                                                image_info.append({
                                                    'location': f'table_{table_idx}_row_{row_idx}_cell_{cell_idx}',
                                                    'question_number': current_question_number,
                                                    'image_type': image_type,
                                                    'text_context': context_text,
                                                    'image_path': image_path
                                                })
                            except Exception as e:
                                # 单个图片提取失败不影响整体
                                self._add_log(f"  ⚠️ 表格单元格图片提取失败: {str(e)}", 'warning')

            return '\n'.join(text_parts), image_info

        except ImportError:
            raise Exception("需要安装 python-docx: pip install python-docx")
        except Exception as e:
            raise Exception(f"提取Word文档失败: {str(e)}")

    def _extract_and_save_image(self, drawing_element, document, image_index: int) -> Optional[str]:
        """
        从drawing元素中提取图片并保存到文件

        Args:
            drawing_element: Word文档中的drawing元素
            document: Document对象，用于访问图片关系
            image_index: 图片索引（用于生成文件名）

        Returns:
            保存的图片文件路径（相对路径），失败则返回None
        """
        try:
            import uuid
            from docx.oxml.ns import qn

            # 使用python-docx的ImagePart方式提取图片（更可靠）
            # 先尝试从drawing元素中查找blip
            blips = drawing_element.xpath('.//a:blip')
            if not blips:
                # 尝试其他可能的XPath
                blips = drawing_element.xpath('.//*[local-name()="blip"]')

            if not blips:
                return None

            blip = blips[0]
            # 获取图片的r:embed属性（指向图片关系ID）
            embed_attr = blip.get(qn('r:embed'))
            if not embed_attr:
                # 尝试不带命名空间的属性
                embed_attr = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if not embed_attr:
                return None

            # 从document part中获取图片关系
            try:
                image_part = document.part.related_parts[embed_attr]
            except KeyError:
                self._add_log(f"  ⚠️ 找不到图片关系: {embed_attr}", 'warning')
                return None

            # 获取图片数据
            image_bytes = image_part.blob

            if not image_bytes:
                self._add_log(f"  ⚠️ 图片数据为空", 'warning')
                return None

            # 检测图片格式
            image_format = image_part.content_type.split('/')[-1].upper()
            if image_format == 'JPEG':
                ext = '.jpg'
            elif image_format == 'PNG':
                ext = '.png'
            else:
                ext = '.png'  # 默认使用png

            # 生成唯一文件名
            unique_id = uuid.uuid4().hex[:8]
            filename = f"image_{unique_id}_{image_index:04d}{ext}"
            file_path = Path(self.upload_folder) / filename

            # 确保目录存在
            file_path.parent.mkdir(parents=True, exist_ok=True)

            # 保存图片
            with open(file_path, 'wb') as f:
                f.write(image_bytes)

            # 转换为相对于uploads目录的相对路径
            # 文件路径应该是 uploads/images/filename 或 uploads/filename 格式
            # 我们需要返回相对于 uploads 根目录的路径
            try:
                # 获取uploads根目录（向上查找直到找到uploads）
                path_parts = Path(file_path).parts
                uploads_index = -1
                for i, part in enumerate(path_parts):
                    if part == 'uploads':
                        uploads_index = i
                        break

                if uploads_index >= 0 and uploads_index < len(path_parts) - 1:
                    # 从uploads之后的所有部分组成相对路径
                    relative_parts = path_parts[uploads_index+1:]
                    relative_path_str = '/'.join(relative_parts)
                else:
                    # 如果找不到uploads目录，只返回文件名
                    relative_path_str = filename

                self._add_log(f"  ✓ 提取图片: {filename} -> {relative_path_str} ({len(image_bytes)} bytes)", 'info')
                self._add_log(f"  🔍 调试: file_path={file_path}, uploads_index={uploads_index}, 返回路径={relative_path_str}", 'debug')
                return relative_path_str
            except Exception as e:
                self._add_log(f"  ⚠️ 计算相对路径失败: {str(e)}, 使用文件名: {filename}", 'warning')
                return filename

        except Exception as e:
            self._add_log(f"  ⚠️ 提取图片失败: {str(e)}", 'warning')
            # 不打印完整堆栈，避免混乱
            return None

    def _get_image_context(self, all_paragraphs: List[str], current_para_idx: int, current_question_number: int) -> str:
        """
        获取图片的上下文文本

        Args:
            all_paragraphs: 所有段落文本列表
            current_para_idx: 当前段落的索引
            current_question_number: 当前题号

        Returns:
            上下文文本
        """
        # 优先向后查找，找到第一个非空的段落，并且要检查是否是选项文本
        for i in range(current_para_idx + 1, min(current_para_idx + 5, len(all_paragraphs))):
            if all_paragraphs[i].strip():
                # 如果这段文本是选项格式（以A-D开头），返回它
                text = all_paragraphs[i].strip()
                if re.match(r'^[A-D][\.、)\s]', text):
                    return text[:100]
                # 否则也返回，让后面的逻辑判断
                return text[:100]

        # 如果向后找不到，向前查找
        for i in range(current_para_idx - 1, max(current_para_idx - 3, -1), -1):
            if all_paragraphs[i].strip():
                return all_paragraphs[i][:100]

        # 如果都找不到，返回当前题号
        if current_question_number:
            return f"第{current_question_number}题"

        return "未知上下文"

    def _determine_image_type(self, text: str, question_number: int) -> str:
        """判断图片类型（题目图片或选项图片）"""
        if not text:
            return 'unknown'

        # 如果文本包含 A、B、C、D 等选项标记，且格式类似选项，则认为是选项图片
        option_patterns = [
            r'^[A-D][\.、]\s*',
            r'^\([A-D]\)\s*',
            r'^[A-D]\s+',
            r'^\s*[A-D]\s*[、.]'
        ]

        for pattern in option_patterns:
            if re.match(pattern, text.strip()):
                return 'option_image'

        # 否则认为是题目图片
        return 'question_image'

    def _parse_with_ai(self, text: str, image_info: List[Dict] = None) -> List[Dict[str, Any]]:
        """
        使用AI从文本中解析试题（一次性将整个文档传给AI解析）

        Args:
            text: 文本内容
            image_info: 图片信息列表

        Returns:
            试题列表
        """
        try:
            # 获取AI服务
            self._add_log("  正在获取AI服务...", 'info', progress=40)
            ai_service = get_ai_service()
            ai_provider = ai_service.__class__.__name__ if hasattr(ai_service, '__class__') else 'Unknown'
            model_name = getattr(ai_service, 'model', 'unknown')
            api_url = getattr(ai_service, 'api_url', '')
            max_tokens_config = getattr(ai_service, 'max_tokens', 2000)
            temperature_config = getattr(ai_service, 'temperature', 0.7)

            # 根据模型名称自动设置合适的max_tokens
            # reasoning/reasoner模型支持更大的输出
            if 'reason' in model_name.lower() or 'r1' in model_name.lower():
                max_tokens = 64000  # DeepSeek R1 reasoning模型支持64K输出
                self._add_log(f"  检测到推理模型，使用max_tokens={max_tokens}", 'info')
            else:
                # chat模型，使用配置中的值，但不超过8192
                max_tokens = min(max_tokens_config, 8192)
                if max_tokens != max_tokens_config:
                    self._add_log(f"  Chat模型限制max_tokens为{max_tokens}（配置值: {max_tokens_config}）", 'info')
                else:
                    self._add_log(f"  使用配置的max_tokens={max_tokens}", 'info')

            # 使用配置的温度
            temperature = temperature_config
            self._add_log(f"  使用配置的温度: {temperature}", 'info')

            # 提取域名显示
            if '://' in api_url:
                api_domain = api_url.split('://')[1].split('/')[0]
            else:
                api_domain = api_url
            self._add_log(f"  AI服务已就绪: {ai_provider} (模型: {model_name})", 'success', progress=45)
            self._add_log(f"  API地址: {api_domain}", 'info', progress=45)

            # 构建提示词
            self._add_log("  正在构建AI提示词...", 'info', progress=50)
            prompt = self._build_ai_prompt(text, image_info)

            # 发送整个文档给AI解析
            self._add_log(f"  正在发送文档给AI（共 {len(text)} 个字符）...", 'info', progress=55)

            messages = [
                {
                    "role": "system",
                    "content": "你是一个专业的试题解析助手，擅长从文档中提取各种类型的试题并转换为标准JSON格式。"
                },
                {"role": "user", "content": prompt}
            ]

            # 调用AI服务，使用自动确定的max_tokens和配置的温度
            response = ai_service.chat(messages, max_tokens=max_tokens, temperature=temperature)

            self._add_log("  AI返回结果，正在解析JSON...", 'info', progress=80)

            # 检查响应内容
            content = response.get('content', '')
            if not content:
                self._add_log("  错误：AI返回的content为空", 'error')
                raise Exception("AI返回的内容为空")

            self._add_log(f"  AI响应内容长度: {len(content)} 字符", 'info')

            # 解析AI返回的JSON
            questions = self._parse_ai_response(content)

            self._add_log(f"  AI解析完成，共获得 {len(questions)} 道试题", 'success', progress=90)

            # 根据实际图片信息更新题目的图片标记
            if image_info and questions:
                questions = self._update_question_image_flags(questions, image_info)
                self._add_log(f"  已根据实际图片信息更新题目图片标记", 'info')

            return questions

        except Exception as e:
            error_msg = f"AI解析失败: {str(e)}"
            self._add_log(error_msg, 'error')
            # AI解析失败，回退到规则解析
            return self._parse_with_rules(text, image_info)

    def _parse_with_rules(self, text: str, image_info: List[Dict] = None) -> List[Dict[str, Any]]:
        """
        使用规则从文本中解析试题

        Args:
            text: 文本内容

        Returns:
            试题列表
        """
        questions = []
        lines = text.split('\n')
        current_question = None
        question_number = 1
        current_section = None  # 当前题型部分
        current_options = []  # 当前问题的选项

        # 题型检测
        section_patterns = {
            'single_choice': [r'一、单选题'],
            'multiple_choice': [r'二、多选题'],
            'judgment': [r'[一二三]、判断题'],
            'fill_blank': [r'四、填空题'],
            'subjective': [r'五、简答题|编程题']
        }

        # 问题开始模式
        question_patterns = [
            r'^(\d+)\.\s*',  # 1.
            r'^(\d+)、\s*',  # 1、
        ]

        # 选项模式
        option_pattern = r'^([A-D])[\.、]\s*(.+)'

        # 答案和解析模式
        answer_pattern = r'正确答案[：:]\s*([A-D]+)'
        explanation_pattern = r'答案解析[：:]\s*'

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 检测题型部分
            section_changed = False
            for q_type, patterns in section_patterns.items():
                for pattern in patterns:
                    if re.search(pattern, line):
                        current_section = q_type
                        section_changed = True
                        break
                if section_changed:
                    break

            if section_changed:
                continue

            # 检查是否是问题开始
            question_match = None
            for pattern in question_patterns:
                match = re.match(pattern, line)
                if match:
                    question_match = match
                    break

            if question_match:
                # 保存上一个问题
                if current_question:
                    # 添加之前收集的选项（仅单选/多选题）
                    if current_question['type'] in ['single_choice', 'multiple_choice']:
                        current_question['options'] = current_options
                    questions.append(current_question)
                    current_options = []

                # 开始新问题
                # 移除题号
                content = re.sub(r'^(\d+)\.|^(\d+)、', '', line).strip()

                # 检查当前题目是否有图片
                has_image = False
                if image_info:
                    for img in image_info:
                        if img['question_number'] == question_number and img['image_type'] == 'question_image':
                            has_image = True
                            break

                current_question = {
                    'type': current_section or 'single_choice',
                    'content': content,
                    'correct_answer': '',
                    'points': 2,  # 默认分值
                    'options': [],
                    'explanation': '',
                    'order_index': question_number,
                    'content_has_image': has_image
                }
                question_number += 1

            elif current_question:
                # 检查是否是选项（仅单选/多选题）
                option_match = re.match(option_pattern, line)
                if option_match and current_question['type'] in ['single_choice', 'multiple_choice']:
                    option_id = option_match.group(1)
                    option_text = option_match.group(2)

                    # 检查选项是否有图片
                    has_image = False
                    if image_info:
                        for img in image_info:
                            if (img['question_number'] == current_question['order_index'] and
                                img['image_type'] == 'option_image' and
                                img['text_context'].strip().startswith(option_id)):
                                has_image = True
                                break

                    current_options.append({
                        'id': option_id,
                        'text': option_text,
                        'has_image': has_image
                    })
                else:
                    # 检查是否是答案
                    answer_match = re.search(answer_pattern, line)
                    if answer_match:
                        current_question['correct_answer'] = answer_match.group(1)
                    # 检查是否是解析
                    elif '答案解析' in line:
                        explanation_text = re.sub(explanation_pattern, '', line).strip()
                        if explanation_text:
                            current_question['explanation'] = explanation_text
                    # 判断题特殊处理：答案格式为"正确答案：正确"或"正确答案：错误"
                    elif current_question['type'] == 'judgment' and '正确答案' in line:
                        answer_text = re.search(r'正确答案[：:]\s*(正确|错误)', line)
                        if answer_text:
                            current_question['correct_answer'] = answer_text.group(1)

        # 保存最后一个问题
        if current_question:
            # 添加之前收集的选项（仅单选/多选题）
            if current_question['type'] in ['single_choice', 'multiple_choice']:
                current_question['options'] = current_options
            questions.append(current_question)

        # 标准化试题格式
        normalized_questions = self._normalize_questions(questions)

        # 根据实际图片信息更新题目图片标记
        if image_info and normalized_questions:
            normalized_questions = self._update_question_image_flags(normalized_questions, image_info)

        return normalized_questions

    def _build_ai_prompt(self, text: str, image_info: List[Dict] = None) -> str:
        """构建AI提示词"""
        # 构建图片信息说明
        image_info_text = ""
        if image_info:
            image_info_text = "\n\n**图片信息：**\n"
            for idx, img in enumerate(image_info):
                img_type_cn = "题目图片" if img['image_type'] == 'question_image' else "选项图片"
                question_num = f"第{img['question_number']}题" if img['question_number'] else "未知题号"
                image_info_text += f"- 图片{idx+1}: {img_type_cn}, 位置: {img['location']}, 对应: {question_num}, 上下文: {img['text_context']}\n"

        prompt = f"""请从以下试卷文档中提取所有试题，并转换为JSON格式。{image_info_text}

**要求：**
1. 识别所有题型：单选题(single_choice)、多选题(multiple_choice)、判断题(true_false)、填空题(fill_blank)、简答题(short_answer)
2. 每道题必须包含以下字段：
   - type: 题型（single_choice/multiple_choice/true_false/fill_blank/short_answer）
   - content: 题目内容（完整文本）
   - options: 选项列表（单选/多选题需要），格式为 [{{"id": "A", "text": "...", "has_image": false}}, ...]
   - correct_answer: 正确答案
   - points: 分值（根据题目难度自动判断，单选题2-5分，多选题3-5分，判断题1-2分，填空题2-3分，简答题5-10分）
   - explanation: 答案解析（从文档中提取）
   - knowledge_point: 考点（分析题目涉及的知识点或考查内容）
   - content_has_image: 题目内容是否包含图片（true/false）
   - order_index: 题号（按顺序从1开始）

3. 对于判断题：正确答案为"true"（正确）或"false"（错误）
4. 对于多选题：正确答案为多个字母组合，如"ABC"
5. 对于填空题和简答题：正确答案为文本内容
6. 图片识别：
   - 如果题目内容中提到"如下图"、"如图所示"、"图片"等关键词，将content_has_image设为true
   - 如果选项中包含图片相关描述，将对应选项的has_image设为true
7. 考点分析：
   - 根据题目内容自动归纳考查的知识点
   - 例如："二叉树的遍历"、"动态规划"、"数组的查找"等

8. **代码格式智能恢复（非常重要）：**
   - 由于从Word文档提取的文本可能丢失了原始代码格式，需要智能识别并恢复
   - 如果遇到空格开头的行（如"    def test():"），这表示代码缩进，必须保留
   - 如果遇到连续的空行，保留它们用于代码块分隔
   - 如果段落包含"def "、"class "、"if "等关键字，可能是代码，保留其格式
   - 对于缩进的代码行，不要删除前导空格
   - 示例：如果文档中有"    print('Hello')"，保留4个空格前缀

9. **代码题目特别处理（非常重要）：**
   - 题目或选项中如果包含代码（Python、Java、C++等），必须完整保留原始格式
   - 保持所有缩进（indentation）不变，特别是Python代码的缩进
   - 保持代码中的换行符和空格，使用\\n表示换行
   - 不要对代码进行格式化或自动修正
   - 如果代码使用制表符或空格缩进，保持原样
   - 代码块中的字符串、注释等内容也要完整保留
   - 示例：如果原文代码是4空格缩进，不要改成2空格或tab

10. **简答题/编程题多要求处理（非常重要）：**
    - 如果题目包含多个要求（如"要求1：... 要求2：..."），在JSON中使用\\n换行符分隔
    - 每个要求另起一行，使用\\n明确标记换行
    - 示例：如果原文是"要求1：写一个函数。要求2：处理边界情况。"
      JSON中应保存为："要求1：写一个函数。\\n要求2：处理边界情况。"
    - 对于多行描述的编程题，保留原文的段落结构，使用\\n换行

11. **JSON格式规范（最关键的要求）：**
    - **返回格式必须是有效的JSON数组，不要包含任何markdown代码块标记（```json或```）**
    - **必须确保JSON完整，不要截断或中断**
    - JSON字符串中的特殊字符必须正确转义：
      - 换行符使用 \\n
      - 制表符使用 \\t
      - 双引号使用 \\"
      - 反斜杠使用 \\\\
    - **不要在JSON字符串中包含未转义的双引号或反斜杠**
    - 代码中的所有空白字符都要保留
    - 不要压缩或删除空行
    - JSON必须使用双引号"，不能使用单引号'
    - 题目和选项中的代码示例、字典示例等需要转义其中的引号
    - **最后一个试题对象后面的括号必须正确闭合，确保数组以 ] 结尾**

**JSON格式示例：**
[
  {{
    "type": "single_choice",
    "content": "以下哪个数据结构遵循先进先出（FIFO）原则？",
    "options": [
      {{"id": "A", "text": "栈", "has_image": false}},
      {{"id": "B", "text": "队列", "has_image": false}},
      {{"id": "C", "text": "链表", "has_image": false}},
      {{"id": "D", "text": "二叉树", "has_image": false}}
    ],
    "correct_answer": "B",
    "points": 2,
    "explanation": "队列遵循先进先出（FIFO）原则，栈遵循后进先出（LIFO）原则。",
    "knowledge_point": "队列的基本特性",
    "content_has_image": false,
    "order_index": 1
  }},
  {{
    "type": "multiple_choice",
    "content": "以下哪些是二叉树遍历的方法？",
    "options": [
      {{"id": "A", "text": "前序遍历", "has_image": false}},
      {{"id": "B", "text": "中序遍历", "has_image": false}},
      {{"id": "C", "text": "后序遍历", "has_image": false}},
      {{"id": "D", "text": "拓扑排序", "has_image": false}}
    ],
    "correct_answer": "ABC",
    "points": 4,
    "explanation": "二叉树的遍历方法包括前序、中序和后序遍历，拓扑排序是图算法。",
    "knowledge_point": "二叉树的遍历",
    "content_has_image": false,
    "order_index": 2
  }}
]

**试卷内容（完整文档，共{len(text)}个字符）：**
{text}

**重要：**
- 只返回JSON格式的试题数组
- 不要包含任何解释性文字、markdown代码块标记或其他额外内容
- 确保所有字段都完整
- 答案必须准确，从文档的"正确答案"或"答案解析"中提取
- 考点要精炼准确，用简短的词组描述，如"队列的基本特性"、"二叉树的遍历"、"动态规划"等
- 代码题必须保持原始缩进和格式，这是最关键的要求
- **返回的必须是纯JSON，不要有任何markdown代码块标记（```json或```）**
- **JSON必须使用双引号"，不能使用单引号'**
- 所有字符串中的双引号必须使用反斜杠转义：\\"
- 题目和选项中的代码示例、字典示例等需要转义其中的引号
- **特别注意：不要返回包含markdown代码块的格式，直接返回JSON数组开头【和结尾】**
|- **确保JSON完整闭合，以 ] 结尾，不要截断**"""

        return prompt

    def _parse_ai_response(self, response: str) -> List[Dict[str, Any]]:
        """
        解析AI响应，提取试题数组

        Args:
            response: AI响应文本

        Returns:
            试题列表
        """
        try:
            # 步骤1: 预处理响应，移除可能的markdown代码块标记
            response = self._fix_json_format(response)
            self._add_log(f"  预处理后响应长度: {len(response)} 字符", 'info')

            # 步骤2: 检查是否以[开头，以]结尾
            if not response.startswith('[') or not response.rstrip().endswith(']'):
                self._add_log("  警告：AI响应不是有效的JSON数组格式", 'warning')
                if not response.rstrip().endswith(']'):
                    self._add_log("  提示：JSON可能被截断，建议减少题目数量或增加max_tokens参数", 'warning')
                    self._add_log(f"  JSON响应前200字符: {response[:200]}", 'info')
                    self._add_log(f"  JSON响应后200字符: {response[-200:]}", 'info')

            # 步骤3: 尝试使用json5解析（更宽松）
            if HAS_JSON5:
                try:
                    questions = json5.loads(response)
                    self._add_log(f"  json5解析成功，获得 {len(questions) if isinstance(questions, list) else 1} 个对象", 'success')
                    return self._normalize_questions(questions)
                except Exception as e:
                    self._add_log(f"  json5解析失败: {str(e)[:200]}", 'warning')
                    self._add_log(f"  错误位置前200字符: {response[:200]}", 'warning')
                    self._add_log(f"  错误位置后200字符: {response[-200:]}", 'warning')

            # 步骤4: 回退到标准JSON解析
            try:
                questions = json.loads(response)
                self._add_log(f"  标准JSON解析成功，获得 {len(questions) if isinstance(questions, list) else 1} 个对象", 'success')
                return self._normalize_questions(questions)
            except json.JSONDecodeError as e:
                error_msg = f"JSON解析失败: {str(e)}"
                self._add_log(error_msg, 'error')
                self._add_log(f"  错误详情: {e.msg} (行 {e.lineno}, 列 {e.colno})", 'error')
                self._add_log(f"  JSON响应前300字符: {response[:300]}", 'error')
                self._add_log(f"  JSON响应后300字符: {response[-300:]}", 'error')
                return []

        except Exception as e:
            error_msg = f"解析AI响应时发生异常: {str(e)}"
            self._add_log(error_msg, 'error')
            return []

    def _fix_json_format(self, json_str: str) -> str:
        """
        修复JSON格式问题 - 只移除markdown代码块标记

        Args:
            json_str: 原始JSON字符串

        Returns:
            修复后的JSON字符串
        """
        # 移除可能的markdown代码块标记（```json 和 ```）
        json_str = re.sub(r'```json\s*', '', json_str)
        json_str = re.sub(r'```\s*', '', json_str)
        json_str = json_str.replace('```', '')

        # 移除控制字符（保留换行符和制表符）
        json_str = ''.join(char for char in json_str if ord(char) >= 32 or char in ['\n', '\t'])

        # 移除开头和结尾的空白字符
        json_str = json_str.strip()

        return json_str

    def _update_question_image_flags(self, questions: List[Dict[str, Any]], image_info: List[Dict]) -> List[Dict[str, Any]]:
        """
        根据实际的图片信息更新题目的图片标记和图片路径

        Args:
            questions: 试题列表
            image_info: 图片信息列表

        Returns:
            更新后的试题列表
        """
        self._add_log(f"  开始根据实际图片信息更新题目图片信息，检测到 {len(image_info)} 个图片", 'info')

        # 按题号分组图片
        images_by_question = {}
        for img in image_info:
            q_num = img['question_number']
            if q_num not in images_by_question:
                images_by_question[q_num] = {'question_images': [], 'option_images': {}}
            if img['image_type'] == 'question_image':
                images_by_question[q_num]['question_images'].append(img)
            elif img['image_type'] == 'option_image':
                # 查找选项ID
                option_id = self._extract_option_id(img['text_context'])
                if option_id:
                    if option_id not in images_by_question[q_num]['option_images']:
                        images_by_question[q_num]['option_images'][option_id] = []
                    images_by_question[q_num]['option_images'][option_id].append(img)
                else:
                    # 如果无法识别选项ID，添加到未知列表，后面按顺序分配
                    if 'unknown_options' not in images_by_question[q_num]:
                        images_by_question[q_num]['unknown_options'] = []
                    images_by_question[q_num]['unknown_options'].append(img)

        for question in questions:
            order_index = question.get('order_index', 0)

            # 检查题目是否有图片
            if order_index in images_by_question:
                question_images = images_by_question[order_index]['question_images']
                if question_images:
                    question['content_has_image'] = True
                    # 使用第一张题目图片的路径
                    question['image_path'] = question_images[0].get('image_path')
                    self._add_log(f"    第{order_index}题包含图片: {question_images[0].get('image_path')}", 'info')
                else:
                    question['content_has_image'] = False

                # 检查选项是否有图片
                if question.get('type') in ['single_choice', 'multiple_choice'] and 'options' in question:
                    option_images = images_by_question[order_index]['option_images']

                    # 获取未知选项图片（无法识别选项ID的）
                    unknown_options = images_by_question[order_index].get('unknown_options', [])

                    # 为每个选项分配图片
                    option_id_order = ['A', 'B', 'C', 'D', 'E', 'F']  # 选项ID顺序
                    unknown_index = 0  # 未知图片的索引

                    for option in question['options']:
                        option_id = option.get('id', '')

                        # 优先使用已识别的选项图片
                        if option_id in option_images and option_images[option_id]:
                            option['has_image'] = True
                            option['image_path'] = option_images[option_id][0].get('image_path')
                            self._add_log(f"    第{order_index}题选项{option_id}包含图片(已识别): {option['image_path']}", 'info')
                        # 如果有未识别的图片，按选项顺序分配
                        elif unknown_options and unknown_index < len(unknown_options):
                            option['has_image'] = True
                            option['image_path'] = unknown_options[unknown_index].get('image_path')
                            self._add_log(f"    第{order_index}题选项{option_id}包含图片(自动分配): {option['image_path']}", 'info')
                            unknown_index += 1
                        else:
                            option['has_image'] = False
            else:
                question['content_has_image'] = False
                if question.get('type') in ['single_choice', 'multiple_choice'] and 'options' in question:
                    for option in question['options']:
                        option['has_image'] = False

        return questions

    def _extract_option_id(self, text: str) -> Optional[str]:
        """
        从文本中提取选项ID（A、B、C、D等）

        Args:
            text: 文本内容

        Returns:
            选项ID，未找到则返回None
        """
        if not text:
            return None

        text = text.strip()

        # 匹配选项ID模式 - 先检查开头
        patterns = [
            r'^([A-D])[\.、]\s*',
            r'^\(([A-D])\)\s*',
            r'^([A-D])\s+',
            r'^\s*选项\s*([A-D])',
            r'^\s*[A-D]\s*[、.]'
        ]

        for pattern in patterns:
            match = re.match(pattern, text)
            if match:
                return match.group(1).upper()

        # 如果开头没有匹配到，尝试在文本中查找（处理像"A. xxx"出现在文本中间的情况）
        match_patterns = [
            r'([A-D])[\.、]\s*',
            r'\(([A-D])\)\s*',
            r'选项\s*([A-D])'
        ]

        for pattern in match_patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1).upper()

        return None

    def _normalize_questions(self, questions: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        标准化试题格式

        Args:
            questions: 原始试题列表

        Returns:
            标准化后的试题列表
        """
        normalized = []

        self._add_log(f"  标准化试题开始，原始题目数量: {len(questions)}", 'info')

        for idx, q in enumerate(questions):
            # 显示原始题目的所有字段和值
            self._add_log(f"  第{idx+1}题原始数据: {str(q)[:200]}", 'info')

            # 验证必需字段
            if not q.get('content'):
                self._add_log(f"  第{idx+1}题被过滤：缺少content字段，题目数据: {list(q.keys())}", 'warning')
                continue

            # 标准化题型
            type_map = {
                '单选题': 'single_choice',
                '多选题': 'multiple_choice',
                '判断题': 'true_false',
                '填空题': 'fill_blank',
                '简答题': 'short_answer',
                'single_choice': 'single_choice',
                'multiple_choice': 'multiple_choice',
                'judgment': 'true_false',
                'subjective': 'short_answer',
                'true_false': 'true_false',
                'fill_blank': 'fill_blank',
                'short_answer': 'short_answer'
            }
            q_type = q.get('type', 'single_choice')
            q_type = type_map.get(q_type, 'single_choice')

            # 标准化选项
            options = q.get('options', [])
            if q_type in ['single_choice', 'multiple_choice']:
                if isinstance(options, list):
                    # 确保选项格式正确
                    normalized_options = []
                    for opt in options:
                        if isinstance(opt, dict):
                            if 'id' not in opt:
                                opt['id'] = chr(65 + len(normalized_options))
                            if 'text' in opt:
                                normalized_options.append(opt)
                    options = normalized_options
                else:
                    options = []

            # 标准化答案
            correct_answer = q.get('correct_answer', '')
            if q_type == 'true_false':
                # 判断题答案标准化为 true/false
                if correct_answer in ['正确', 'true', 'True', '是', 'A']:
                    correct_answer = 'true'
                elif correct_answer in ['错误', 'false', 'False', '否', 'B']:
                    correct_answer = 'false'

            # 安全处理分值
            try:
                points = int(q.get('points', 2))
            except (ValueError, TypeError):
                points = 2

            normalized_q = {
                'type': q_type,
                'content': q['content'],
                'correct_answer': str(correct_answer),
                'points': points,
                'options': options,
                'explanation': q.get('explanation', ''),
                'knowledge_point': q.get('knowledge_point', ''),
                'order_index': idx + 1
            }

            normalized.append(normalized_q)

        self._add_log(f"  标准化完成，有效题目数量: {len(normalized)}", 'info')
        return normalized

    def parse_to_exam(self, file_path: str, exam_title: str,
                     subject_id: int, level_id: int) -> Dict[str, Any]:
        """
        解析文档并转换为考试数据

        Args:
            file_path: 文档路径
            exam_title: 考试标题
            subject_id: 科目ID
            level_id: 等级ID

        Returns:
            考试数据字典
        """
        result = self.parse_document(file_path)

        questions = result.get('questions', [])
        total_points = sum(q['points'] for q in questions)

        return {
            'title': exam_title,
            'subject_id': subject_id,
            'level_id': level_id,
            'description': f"从文档 '{Path(file_path).name}' 导入",
            'duration_minutes': len(questions) * 2 if questions else 60,
            'total_points': total_points,
            'questions': questions,
            'parse_method': result.get('parse_method', 'ai'),
            'question_count': len(questions)
        }
