"""DOCX 文档文本和图片提取器"""

import re
import uuid
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path


class DocxExtractor:
    """DOCX 文档文本和图片提取器"""

    def __init__(self, upload_folder: str = 'uploads/images', extract_images: bool = True):
        """
        初始化提取器

        Args:
            upload_folder: 图片保存目录
            extract_images: 是否提取图片
        """
        self.upload_folder = upload_folder
        self.extract_images = extract_images
        if extract_images:
            Path(upload_folder).mkdir(parents=True, exist_ok=True)

    def extract_text(self, file_path: str) -> str:
        """只提取文本，不提取图片"""
        try:
            from docx import Document
            doc = Document(file_path)

            text_parts = []
            for para in doc.paragraphs:
                text_parts.append(para.text)

            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())

            return '\n'.join(text_parts)

        except ImportError:
            raise Exception("需要安装 python-docx: pip install python-docx")
        except Exception as e:
            raise Exception(f"提取文本失败: {str(e)}")

    def extract_text_and_images(self, file_path: str) -> Tuple[str, List[Dict]]:
        """提取文本和图片"""
        try:
            from docx import Document
            doc = Document(file_path)

            text_parts = []
            image_info = []
            current_question_number = None
            line_number = 0
            global_image_index = 0

            # 首先收集所有段落文本
            all_paragraphs_text = []
            for para in doc.paragraphs:
                all_paragraphs_text.append(para.text.strip())

            # 提取段落内容和图片
            for para_idx, para in enumerate(doc.paragraphs):
                line_number += 1
                para_text = para.text
                stripped_text = para_text.strip()

                if para_text or (para_idx > 0 and para.text == '' and len(all_paragraphs_text) > 0):
                    if stripped_text:
                        # 检测题号
                        match = re.match(r'^(\d+)[\.、]\s*', stripped_text)
                        if match:
                            current_question_number = int(match.group(1))
                        text_parts.append(para_text)
                    else:
                        text_parts.append('')

                # 提取段落中的图片
                if self.extract_images:
                    try:
                        for run in para.runs:
                            for drawing in run._element.xpath('.//w:drawing'):
                                pics = drawing.xpath('.//pic:pic')
                                if pics:
                                    context_text = self._get_image_context(all_paragraphs_text, para_idx, current_question_number)
                                    image_type = self._determine_image_type(context_text, current_question_number)
                                    image_path = self._extract_and_save_image(drawing, doc, global_image_index)
                                    global_image_index += 1

                                    image_info.append({
                                        'location': f'paragraph_{line_number}',
                                        'question_number': current_question_number,
                                        'image_type': image_type,
                                        'text_context': context_text,
                                        'image_path': image_path
                                    })
                    except Exception:
                        pass

            # 提取表格内容
            table_line_number = line_number
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            text_parts.append(cell_text)
                            table_line_number += 1

                        # 检查单元格中的图片
                        if self.extract_images:
                            try:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        for drawing in run._element.xpath('.//w:drawing'):
                                            pics = drawing.xpath('.//pic:pic')
                                            if pics:
                                                context_text = cell.text.strip()
                                                image_type = self._determine_image_type(context_text, current_question_number)
                                                image_path = self._extract_and_save_image(drawing, doc, global_image_index)
                                                global_image_index += 1

                                                image_info.append({
                                                    'location': f'table_{table_idx}_row_{row_idx}_cell_{cell_idx}',
                                                    'question_number': current_question_number,
                                                    'image_type': image_type,
                                                    'text_context': context_text,
                                                    'image_path': image_path
                                                })
                            except Exception:
                                pass

            return '\n'.join(text_parts), image_info

        except ImportError:
            raise Exception("需要安装 python-docx: pip install python-docx")
        except Exception as e:
            raise Exception(f"提取Word文档失败: {str(e)}")

    def _extract_and_save_image(self, drawing_element, document, image_index: int) -> Optional[str]:
        """从drawing元素中提取图片并保存"""
        try:
            from docx.oxml.ns import qn

            blips = drawing_element.xpath('.//a:blip')
            if not blips:
                blips = drawing_element.xpath('.//*[local-name()="blip"]')

            if not blips:
                return None

            blip = blips[0]
            embed_attr = blip.get(qn('r:embed'))
            if not embed_attr:
                embed_attr = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if not embed_attr:
                return None

            try:
                image_part = document.part.related_parts[embed_attr]
            except KeyError:
                return None

            image_bytes = image_part.blob
            if not image_bytes:
                return None

            image_format = image_part.content_type.split('/')[-1].upper()
            ext = '.jpg' if image_format == 'JPEG' else '.png'

            unique_id = uuid.uuid4().hex[:8]
            filename = f"image_{unique_id}_{image_index:04d}{ext}"
            file_path = Path(self.upload_folder) / filename
            file_path.parent.mkdir(parents=True, exist_ok=True)

            with open(file_path, 'wb') as f:
                f.write(image_bytes)

            # 返回相对路径
            path_parts = Path(file_path).parts
            uploads_index = -1
            for i, part in enumerate(path_parts):
                if part == 'uploads':
                    uploads_index = i
                    break

            if uploads_index >= 0 and uploads_index < len(path_parts) - 1:
                relative_parts = path_parts[uploads_index+1:]
                return '/'.join(relative_parts)
            else:
                return filename

        except Exception:
            return None

    def _get_image_context(self, all_paragraphs: List[str], current_para_idx: int, current_question_number: int) -> str:
        """获取图片的上下文文本"""
        if current_para_idx < len(all_paragraphs):
            current_text = all_paragraphs[current_para_idx].strip()
            if current_text:
                match = re.match(r'^([A-D])[\.、)\s]', current_text)
                if match:
                    return current_text[:100]

        for i in range(current_para_idx - 1, max(current_para_idx - 5, -1), -1):
            if all_paragraphs[i].strip():
                text = all_paragraphs[i].strip()
                if re.match(r'^[A-D][\.、)\s]', text):
                    return text[:100]
                return text[:100]

        for i in range(current_para_idx + 1, min(current_para_idx + 5, len(all_paragraphs))):
            if all_paragraphs[i].strip():
                return all_paragraphs[i][:100]

        if current_question_number:
            return f"第{current_question_number}题"

        return "未知上下文"

    def _determine_image_type(self, text: str, question_number: int) -> str:
        """判断图片类型（题目图片或选项图片）"""
        if not text:
            return 'unknown'

        option_patterns = [
            r'^[A-D][\.、]\s*',
            r'^\([A-D]\)\s*',
            r'^[A-D]\s+',
            r'^\s*[A-D]\s*[、.]'
        ]

        for pattern in option_patterns:
            if re.match(pattern, text.strip()):
                return 'option_image'

        return 'question_image'
